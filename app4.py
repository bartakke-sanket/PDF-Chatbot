import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import pdfplumber
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import os
from PIL import Image
import io

# Load environment variables
load_dotenv()

# Configure Google Generative AI API Key
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Function to read text from PDFs using pdfplumber (and handle tables/images)
def extract_from_pdf(pdf_docs):
    text = ''
    tables = []
    images = []

    for pdf in pdf_docs:
        with pdfplumber.open(pdf) as pdf_reader:
            for page in pdf_reader.pages:
                text += page.extract_text() or ''
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
                for image in page.images:
                    image_data = {
                        "x0": image["x0"],
                        "y0": image["y0"],
                        "x1": image["x1"],
                        "y1": image["y1"],
                        "width": image["width"],
                        "height": image["height"],
                        "name": image.get("name", None)
                    }
                    images.append(image_data)

    return text, tables, images

# Function to read text from .docx files (and handle tables/images)
def extract_from_docx(docx_docs):
    text = ''
    tables = []
    images = []

    for docx in docx_docs:
        doc = Document(docx)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image = rel.target_part.blob
                images.append(image)

    return text, tables, images

# Function to split text into chunks
def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

# Function to create a vector store for text
def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model='models/embedding-001')
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local('faiss_index')

# Function to get the conversational chain (using the Generative AI model)
def get_conversational_chain():
    prompt_template = """
    You are chatting with a document that may contain text, tables, and images.
    Answer questions based on the provided context. For tables, respond in tabular format.
    If an image is asked about, provide the image.
    If the answer is not in the context, say "Answer not available in the provided context."
    Do not provide incorrect answers.\n\n
    Context:\n{context}?\n
    Question:\n{question}\n
    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain

# Function to handle user queries
# Function to handle user queries, including image responses
def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()
    response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)

    # Check if the response mentions an image and if images exist in the session state
    if "image" in user_question.lower() and "extracted_images" in st.session_state:
        # Assume that user is asking for an image, so fetch the stored images
        images = st.session_state.extracted_images
        if images:
            st.write("### Here are the extracted images related to your query:")
            for img in images:
                try:
                    img_pil = Image.open(io.BytesIO(img))
                    st.image(img_pil, caption="Extracted Image")
                except Exception as e:
                    st.write(f"Error displaying image: {e}")
        else:
            return "No images found in the context."

    return response["output_text"]

# def user_input(user_question):
#     embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#     new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
#     docs = new_db.similarity_search(user_question)

#     chain = get_conversational_chain()
#     response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)

#     return response["output_text"]

# Function to display tables in a clean format
def display_tables(tables):
    if tables:
        st.write("### Extracted Tables:")
        for table in tables:
            st.table(table)
    else:
        st.write("No tables found.")

# Function to display images
def display_images(images):
    if images:
        st.write("### Extracted Images:")
        for img in images:
            if isinstance(img, bytes):  # Check if img is image data (bytes)
                try:
                    img_pil = Image.open(io.BytesIO(img))
                    st.image(img_pil, caption="Extracted Image")
                except Exception as e:
                    st.write(f"Error displaying image: {e}")
            else:
                st.write(f"Image metadata: {img}")  # Display metadata info if it's a dict
    else:
        st.write("No images found.")

# Main function to create Streamlit app interface
def main():
    st.set_page_config(page_title="Chat with Multiple Documents", page_icon=":books:")
    st.title("Chat with Your Documents :books:")

    # Initialize session state for storing chat history and images
    if "conversation" not in st.session_state:
        st.session_state.conversation = []
    if "extracted_images" not in st.session_state:
        st.session_state.extracted_images = []

    # Sidebar for uploading PDFs and DOCX files
    with st.sidebar:
        st.subheader("Upload Your Documents (PDFs or DOCX)")
        docs = st.file_uploader("Upload PDFs or Word Documents", type=["pdf", "docx"], accept_multiple_files=True)

        if st.button("Process"):
            if docs:
                with st.spinner("Processing your documents..."):
                    # Separate PDFs and DOCX files
                    pdf_docs = [doc for doc in docs if doc.type == "application/pdf"]
                    docx_docs = [doc for doc in docs if doc.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]

                    # Extract text, tables, and images from PDFs and DOCX
                    pdf_text, pdf_tables, pdf_images = extract_from_pdf(pdf_docs) if pdf_docs else ("", [], [])
                    docx_text, docx_tables, docx_images = extract_from_docx(docx_docs) if docx_docs else ("", [], [])

                    # Combine text from both sources
                    combined_text = pdf_text + docx_text

                    # Store the extracted images in session state
                    st.session_state.extracted_images = pdf_images + docx_images

                    # Split text into chunks and create vector store
                    text_chunks = get_text_chunks(combined_text)
                    get_vector_store(text_chunks)

                    st.success("Documents processed successfully!")

                    # Display tables and images
                    display_tables(pdf_tables + docx_tables)
                    display_images(pdf_images + docx_images)
            else:
                st.error("Please upload PDF or DOCX files before processing.")

    # Main Chat Interface
    st.subheader("Chat with Your Documents")

    user_question = st.text_input("Enter your question here:")

    if st.button("Send"):
        if user_question:
            with st.spinner("Fetching response..."):
                response = user_input(user_question)
                # Append the query and response to the session state conversation
                st.session_state.conversation.append({"question": user_question, "response": response})

    # Display conversation history
    if st.session_state.conversation:
        st.write("### Chat History:")
        for chat in st.session_state.conversation:
            st.write(f"**You**: {chat['question']}")
            st.write(f"**Bot**: {chat['response']}")


# Run the Streamlit app
if __name__ == "__main__":
    main()
